from docx import Document as DocxDocument

from app.services.extraction.docx_extractor import extract_docx_text


def test_extract_text_from_paragraphs(tmp_path):
    docx_path = tmp_path / "test.docx"
    doc = DocxDocument()
    doc.add_paragraph("First paragraph of content.")
    doc.add_paragraph("Second paragraph of content.")
    doc.save(str(docx_path))

    result = extract_docx_text(str(docx_path))

    assert "First paragraph of content." in result
    assert "Second paragraph of content." in result


def test_extract_text_from_table(tmp_path):
    docx_path = tmp_path / "table.docx"
    doc = DocxDocument()
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Name"
    table.cell(0, 1).text = "Role"
    table.cell(1, 0).text = "Andualem"
    table.cell(1, 1).text = "Backend Developer"
    doc.save(str(docx_path))

    result = extract_docx_text(str(docx_path))

    assert "Name | Role" in result
    assert "Andualem | Backend Developer" in result


def test_extract_empty_docx_returns_empty_string(tmp_path):
    docx_path = tmp_path / "empty.docx"
    doc = DocxDocument()
    doc.save(str(docx_path))

    result = extract_docx_text(str(docx_path))

    assert result == ""