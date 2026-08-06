from docx import Document as DocxDocument


def extract_docx_text(file_path: str) -> str:
    doc = DocxDocument(file_path)

    parts = []
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)

    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells)
            if row_text.strip(" |"):
                parts.append(row_text)

    return "\n".join(parts).strip()